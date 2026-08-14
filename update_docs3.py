import glob
try:
    from docx import Document
except ImportError as e:
    raise SystemExit("Missing dependency: python-docx (pip install python-docx)") from e

def process_text(text):
    text = text.replace("infrastructure pivot thesis", "agentic infrastructure pivot thesis")
    text = text.replace("operational infrastructure", "agentic infrastructure")
    text = text.replace("digital infrastructure", "agentic infrastructure")
    text = text.replace("Digital infrastructure", "Agentic infrastructure")
    text = text.replace("structural digital infrastructure", "structural agentic infrastructure")
    text = text.replace("infrastructure layer", "agentic infrastructure layer")
    return text

for ext in ['**/*.md', '**/*.html']:
    for path in glob.glob(ext, recursive=True):
        if 'node_modules' in path or '.git' in path: continue
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        new_content = process_text(content)
        
        if content != new_content:
            print(f"Updated {path}")
            with open(path, 'w', encoding='utf-8') as f:
                f.write(new_content)

for path in glob.glob('**/*.docx', recursive=True):
    if 'node_modules' in path or '.git' in path: continue
    try:
        doc = Document(path)
        changed = False
        for p in doc.paragraphs:
            new_text = process_text(p.text)
            if new_text != p.text:
                p.text = new_text
                changed = True
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        new_text = process_text(p.text)
                        if new_text != p.text:
                            p.text = new_text
                            changed = True
                            
        if changed:
            print(f"Updated {path}")
            doc.save(path)
    except Exception as e:
        print(f"Error processing {path}: {e}")

