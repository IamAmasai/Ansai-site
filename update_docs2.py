import os
import glob
import re
from docx import Document

def process_text(text):
    text = text.replace("Ansai is an AI agent company. Ansai builds", "Ansai is an AI and agent-native company building")
    text = text.replace("AI is not the product.", "AI and agents are the core product.")
    
    text = text.replace("Why not \"AI agent company\"", "Why \"AI agent company\"")
    text = text.replace("leading with \"AI agent\" invites a comparison set Ansai cannot win", "leading with \"AI agent\" is exactly our identity and what we do")
    text = text.replace("Not \"AI agent company\"", "We are an \"AI agent company\"")
    
    text = text.replace("Nothing on the page describes Ansai as an \"AI agent company\"", "Ensure the page describes Ansai as an \"AI agent company\"")
    text = text.replace("deliberately refusing the default \"AI agent company\" category in favor of", "deliberately adopting the \"AI agent company\" category as")
    
    text = text.replace("What AI is | one tenant / enabler among others, that **permeates** every layer rather than staying in its own lane; named plainly only in technical depth | the leading category word, the headline | Locked", "What AI is | the core of our agentic infrastructure, central to our platform | a mere side feature | Locked")

    text = text.replace("AI is an enabler, not a destination", "AI is both the enabler and the destination")
    
    text = text.replace("agentic infra, AI agent company, data structure company", "data structure company")
    
    # Check if "agentic infrastructure" is replaced anywhere else where it shouldn't be? No, that's fine.
    
    return text

for ext in ['**/*.md', '**/*.html']:
    for path in glob.glob(ext, recursive=True):
        if 'node_modules' in path or '.git' in path: continue
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        new_content = process_text(content)
        
        if content != new_content:
            print(f"Updated {path}")
            with open(path, 'w', encoding='utf-8') as f:
                f.write(new_content)

for path in glob.glob('**/*.docx', recursive=True):
    if 'node_modules' in path or '.git' in path: continue
    try:
        doc = Document(path)
        changed = False
        for p in doc.paragraphs:
            new_text = process_text(p.text)
            if new_text != p.text:
                p.text = new_text
                changed = True
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        new_text = process_text(p.text)
                        if new_text != p.text:
                            p.text = new_text
                            changed = True
                            
        if changed:
            print(f"Updated {path}")
            doc.save(path)
    except Exception as e:
        print(f"Error processing {path}: {e}")

